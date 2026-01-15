import os
import io
import uuid
import base64
import subprocess
import zipfile
import xml.etree.ElementTree as ET
from flask import Flask, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

app = Flask(__name__)

# --- AYARLAR ---
LOGO_URL = "https://static.wixstatic.com/media/06f423_9d350d42007948448e351781e43950c1~mv2.png"

# --- YARDIMCI FONKSƒ∞YONLAR ---
def clean_tag(tag):
    if '}' in tag: return tag.split('}', 1)[1]
    return tag

def clean_attribs(attrib):
    return {clean_tag(k).lower(): v for k, v in attrib.items()}

def get_image_from_data(base64_string):
    try:
        clean_data = base64_string.replace('\n', '').replace('\r', '').replace(' ', '')
        if len(clean_data) < 200: return None
        img_bytes = base64.b64decode(clean_data)
        img_io = io.BytesIO(img_bytes)
        img = Image.open(img_io)
        img.verify()
        img_io.seek(0)
        return img_io
    except:
        return None

def apply_formatting(run, attrib):
    run.font.color.rgb = RGBColor(0, 0, 0)
    if attrib.get('bold') == 'true' or attrib.get('b') == 'true': run.bold = True
    if attrib.get('italic') == 'true' or attrib.get('i') == 'true': run.italic = True
    if attrib.get('underline') == 'true' or attrib.get('u') == 'true': run.underline = True
    f_size = attrib.get('size') or attrib.get('fontsize')
    if f_size and f_size.isdigit():
        try: run.font.size = Pt(int(f_size))
        except: pass

def generate_word_doc(udf_dosya_objesi):
    try:
        with zipfile.ZipFile(udf_dosya_objesi) as z:
            if 'content.xml' not in z.namelist(): return None, "content.xml bulunamadƒ±"
            
            xml_content = z.read('content.xml').decode('utf-8', errors='ignore')
            root = ET.fromstring(xml_content)
            doc = Document()

            global_text = ""
            max_len = 0
            for elem in root.iter():
                if elem.text:
                    l = len(elem.text)
                    if l > max_len: max_len = l; global_text = elem.text
            
            if not global_text: doc.add_paragraph("[Metin i√ßeriƒüi bo≈ü]")

            elements_processed = False
            found_images = [] 
            elements_node = None
            
            for elem in root.iter():
                if 'elements' in clean_tag(elem.tag): elements_node = elem; break
            
            if elements_node is not None and global_text:
                current_p = doc.add_paragraph()
                for item in elements_node.iter():
                    attrib = clean_attribs(item.attrib)
                    tag = clean_tag(item.tag)

                    if tag == 'content' and 'startoffset' in attrib:
                        elements_processed = True
                        try:
                            start = int(attrib['startoffset'])
                            length = int(attrib['length'])
                            chunk = global_text[start : start + length]
                            if '\n' in chunk:
                                lines = chunk.split('\n')
                                for i, line in enumerate(lines):
                                    if line:
                                        run = current_p.add_run(line)
                                        apply_formatting(run, attrib)
                                    if i < len(lines) - 1: current_p = doc.add_paragraph()
                            else:
                                run = current_p.add_run(chunk)
                                apply_formatting(run, attrib)
                        except: pass
                    
                    elif 'imagedata' in attrib:
                        v = attrib['imagedata']
                        if len(v) > 500:
                            img_obj = get_image_from_data(v)
                            if img_obj: found_images.append(img_obj)

            if not elements_processed and global_text:
                p = doc.add_paragraph()
                run = p.add_run(global_text)
                run.font.color.rgb = RGBColor(0, 0, 0)
                for elem in root.iter():
                    for k, v in elem.attrib.items():
                        if len(v) > 500:
                            img_obj = get_image_from_data(v)
                            if img_obj: found_images.append(img_obj)

            if found_images:
                doc.add_page_break()
                baslik = doc.add_heading('DOSYA ƒ∞√áƒ∞NDEKƒ∞ G√ñRSELLER', level=1)
                baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, img_obj in enumerate(found_images):
                    doc.add_paragraph("\n")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    try: run.add_picture(img_obj, width=Inches(4.5))
                    except: continue
                    lbl = doc.add_paragraph(f"#G√∂rsel {i+1}")
                    lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    lbl.style.font.bold = True
                    lbl.style.font.color.rgb = RGBColor(0, 0, 0)

            if 'sign.sgn' in z.namelist():
                doc.add_paragraph("\n\n")
                p = doc.add_paragraph("e-imzalƒ±dƒ±r")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                p.runs[0].font.size = Pt(14)
                p2 = doc.add_paragraph("(Bu belge 5070 sayƒ±lƒ± Kanun uyarƒ±nca imzalanmƒ±≈ütƒ±r)")
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.style.font.size = Pt(9)
                p2.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            return doc, None
    except Exception as e:
        return None, str(e)

# --- ARAY√úZ ---
@app.route('/')
def anasayfa():
    return f'''
    <!doctype html>
    <html lang="tr">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>UDF D√∂n√º≈üt√ºr√ºc√º</title>
        <style>
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background-color: transparent; display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; flex-direction: column; }}
            .card {{ background: #ffffff; padding: 40px; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.15); text-align: center; width: 450px; max-width: 90%; position: relative; }}
            .logo-img {{ max-width: 200px; margin-bottom: 20px; }}
            
            /* Dosya Y√ºkleme */
            .file-upload {{ position: relative; display: inline-block; width: 100%; margin-bottom: 25px; }}
            input[type=file] {{ border: 2px dashed #ccc; padding: 20px; width: 88%; border-radius: 8px; background: #f0f0f0; color: #555; cursor: pointer; transition: 0.3s; }}
            input[type=file]:hover {{ border-color: #999; background: #e9e9e9; }}
            
            /* Butonlar */
            .btn-group {{ display: flex; gap: 15px; margin-bottom: 20px; }}
            .btn {{ flex: 1; padding: 14px; border: none; border-radius: 6px; cursor: pointer; font-size: 15px; font-weight: 600; color: white; transition: transform 0.1s, opacity 0.2s; }}
            .btn:active {{ transform: scale(0.98); }}
            .btn-word {{ background-color: #2b5797; }}
            .btn-pdf {{ background-color: #d32f2f; }}
            .btn:hover {{ opacity: 0.9; }}
            .btn:disabled {{ background-color: #999; cursor: not-allowed; }}

            /* Progress Bar */
            .progress-container {{ display: none; width: 100%; background-color: #d0d0d0; border-radius: 4px; margin: 20px 0; overflow: hidden; }}
            .progress-bar {{ width: 0%; height: 10px; background-color: #666; transition: width 0.4s ease; }}
            .loader-text {{ font-size: 13px; color: #777; margin-top: 5px; display: none; }}
            
            /* Sonu√ß Alanƒ± */
            .result-area {{ display: none; margin-top: 20px; padding: 20px; background: #e0e0e0; border-radius: 8px; border: 1px solid #ccc; }}
            .success-msg {{ color: #333; font-weight: bold; font-size: 18px; margin-bottom: 15px; display: block; }}
            .download-link {{ display: inline-block; text-decoration: none; background: #333; color: white; padding: 12px 25px; border-radius: 5px; font-weight: bold; transition: 0.2s; }}
            .download-link:hover {{ background: #000; }}
            .reset-link {{ cursor:pointer; color:#666; text-decoration:underline; margin-top: 15px; display: inline-block; }}

            /* YENƒ∞: Uygulama Y√ºkle Butonu */
            .install-trigger {{ margin-top: 30px; background: transparent; border: 1px solid #ddd; color: #666; padding: 10px 20px; border-radius: 25px; font-size: 13px; cursor: pointer; display: inline-flex; align-items: center; gap: 8px; transition: 0.2s; }}
            .install-trigger:hover {{ background: #f5f5f5; color: #333; border-color: #bbb; }}
            
            /* Modal Penceresi */
            .modal-overlay {{ display: none; position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: rgba(0,0,0,0.6); z-index: 1000; align-items: center; justify-content: center; backdrop-filter: blur(2px); }}
            .modal {{ background: white; padding: 30px; border-radius: 15px; width: 320px; text-align: center; position: relative; box-shadow: 0 20px 40px rgba(0,0,0,0.2); animation: popUp 0.3s ease; }}
            @keyframes popUp {{ from {{transform: scale(0.8); opacity: 0;}} to {{transform: scale(1); opacity: 1;}} }}
            
            .modal h3 {{ margin-top: 0; color: #222; font-size: 18px; margin-bottom: 15px; }}
            .modal p {{ font-size: 14px; color: #555; line-height: 1.6; margin-bottom: 20px; }}
            .modal-icon {{ font-size: 40px; margin-bottom: 15px; display: block; }}
            
            .modal-close {{ position: absolute; top: 15px; right: 15px; font-size: 24px; cursor: pointer; color: #aaa; line-height: 1; }}
            .modal-close:hover {{ color: #333; }}
            
            .modal-btn {{ background: #333; color: white; border: none; padding: 10px 20px; border-radius: 8px; cursor: pointer; font-size: 14px; }}
        </style>
    </head>
    <body>
        <div class="card">
            <img src="{LOGO_URL}" alt="Logo" class="logo-img">
            <form id="uploadForm">
                <div class="file-upload"><input type="file" id="fileInput" name="dosya" required accept=".udf"></div>
                
                <div class="btn-group" id="btnGroup">
                    <button type="button" class="btn btn-word" onclick="startConversion('word')">D√∂n√º≈üt√ºr: WORD</button>
                    <button type="button" class="btn btn-pdf" onclick="startConversion('pdf')">D√∂n√º≈üt√ºr: PDF</button>
                </div>

                <div class="progress-container" id="progressContainer"><div class="progress-bar" id="progressBar"></div></div>
                <div class="loader-text" id="loaderText">Dosya i≈üleniyor, l√ºtfen bekleyin...</div>

                <div class="result-area" id="resultArea">
                    <span class="success-msg">D√∂n√º≈üt√ºrme Tamamlandƒ±!</span>
                    <a href="#" id="downloadBtn" class="download-link">Dosyayƒ± ƒ∞ndir</a><br>
                    <small class="reset-link" onclick="resetForm()">Yeni Dosya √áevir</small>
                </div>
            </form>

            <button class="install-trigger" onclick="showInstallGuide()">
                üì≤ Uygulamayƒ± Ana Ekrana Ekle
            </button>
        </div>

        <div class="modal-overlay" id="installModal">
            <div class="modal">
                <span class="modal-close" onclick="closeModal()">&times;</span>
                <span class="modal-icon" id="modalIcon">üì±</span>
                <h3 id="modalTitle">Nasƒ±l Eklenir?</h3>
                <p id="modalDesc">Tarayƒ±cƒ± ayarlarƒ±ndan ekleyebilirsiniz.</p>
                <button class="modal-btn" onclick="closeModal()">Anladƒ±m</button>
            </div>
        </div>

        <script>
            // --- CONVERSION LOGIC ---
            function startConversion(type) {{
                var fileInput = document.getElementById('fileInput');
                if (fileInput.files.length === 0) {{ alert("L√ºtfen √∂nce bir UDF dosyasƒ± se√ßin."); return; }}
                document.getElementById('btnGroup').style.display = 'none';
                document.getElementById('progressContainer').style.display = 'block';
                document.getElementById('loaderText').style.display = 'block';
                var progressBar = document.getElementById('progressBar');
                var width = 0;
                var interval = setInterval(function() {{ if (width >= 90) clearInterval(interval); else {{ width++; progressBar.style.width = width + '%'; }} }}, 50);
                var formData = new FormData();
                formData.append('dosya', fileInput.files[0]);
                var url = type === 'word' ? '/indir_word' : '/indir_pdf';
                fetch(url, {{ method: 'POST', body: formData }})
                .then(response => {{ if (response.status !== 200) throw new Error("D√∂n√º≈üt√ºrme hatasƒ±"); return response.blob(); }})
                .then(blob => {{
                    clearInterval(interval); progressBar.style.width = '100%';
                    var downloadUrl = window.URL.createObjectURL(blob);
                    var downloadBtn = document.getElementById('downloadBtn');
                    downloadBtn.href = downloadUrl;
                    downloadBtn.download = fileInput.files[0].name.replace('.udf', type === 'word' ? '.docx' : '.pdf');
                    downloadBtn.innerText = type === 'word' ? "Word Dosyasƒ±nƒ± ƒ∞ndir" : "PDF Dosyasƒ±nƒ± ƒ∞ndir";
                    setTimeout(function() {{
                        document.getElementById('progressContainer').style.display = 'none';
                        document.getElementById('loaderText').style.display = 'none';
                        document.getElementById('resultArea').style.display = 'block';
                    }}, 600);
                }})
                .catch(error => {{ alert("Bir hata olu≈ütu: " + error); resetForm(); }});
            }}
            function resetForm() {{
                document.getElementById('resultArea').style.display = 'none';
                document.getElementById('btnGroup').style.display = 'flex';
                document.getElementById('progressContainer').style.display = 'none';
                document.getElementById('loaderText').style.display = 'none';
                document.getElementById('progressBar').style.width = '0%';
                document.getElementById('fileInput').value = "";
            }}

            // --- INSTALL GUIDE LOGIC ---
            function showInstallGuide() {{
                var modal = document.getElementById('installModal');
                var title = document.getElementById('modalTitle');
                var desc = document.getElementById('modalDesc');
                var icon = document.getElementById('modalIcon');
                var userAgent = navigator.userAgent || navigator.vendor || window.opera;

                modal.style.display = 'flex';

                if (/iPad|iPhone|iPod/.test(userAgent) && !window.MSStream) {{
                    // iOS
                    icon.innerHTML = "üçè";
                    title.innerText = "iPhone'a Ekleme";
                    desc.innerHTML = "1. Tarayƒ±cƒ±nƒ±n en altƒ±ndaki <b>'Payla≈ü'</b> simgesine dokunun.<br><br>2. A≈üaƒüƒ± kaydƒ±rƒ±p <b>'Ana Ekrana Ekle'</b> se√ßeneƒüine basƒ±n.";
                }} else if (/android/i.test(userAgent)) {{
                    // Android
                    icon.innerHTML = "ü§ñ";
                    title.innerText = "Android'e Ekleme";
                    desc.innerHTML = "1. Tarayƒ±cƒ±nƒ±n saƒü √ºst√ºndeki <b>'√ú√ß Nokta'</b> men√ºs√ºne dokunun.<br><br>2. <b>'Uygulamayƒ± Y√ºkle'</b> veya <b>'Ana Ekrana Ekle'</b> se√ßeneƒüine basƒ±n.";
                }} else {{
                    // Desktop
                    icon.innerHTML = "üíª";
                    title.innerText = "Masa√ºst√ºne Ekleme";
                    desc.innerHTML = "Tarayƒ±cƒ±nƒ±zƒ±n men√ºs√ºnden (Genelde adres √ßubuƒüunun saƒüƒ±nda) <b>'Kƒ±sayol Olu≈ütur'</b> veya <b>'Uygulamayƒ± Y√ºkle'</b> diyerek masa√ºst√ºn√ºze ekleyebilirsiniz.";
                }}
            }}

            function closeModal() {{
                document.getElementById('installModal').style.display = 'none';
            }}
            
            window.onclick = function(event) {{
                if (event.target == document.getElementById('installModal')) {{
                    closeModal();
                }}
            }}
        </script>
    </body>
    </html>
    '''

@app.route('/indir_word', methods=['POST'])
def indir_word():
    if 'dosya' not in request.files: return 'Hata', 400
    dosya = request.files['dosya']
    doc, hata = generate_word_doc(dosya)
    if hata: return str(hata), 500
    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    return send_file(mem_file, as_attachment=True, download_name='converted.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/indir_pdf', methods=['POST'])
def indir_pdf():
    if 'dosya' not in request.files: return 'Hata', 400
    dosya = request.files['dosya']
    doc, hata = generate_word_doc(dosya)
    if hata: return str(hata), 500
    unique_id = str(uuid.uuid4())
    temp_docx = f"temp_{unique_id}.docx"
    temp_pdf = f"temp_{unique_id}.pdf"
    try:
        doc.save(temp_docx)
        subprocess.run(['libreoffice', '--headless', '--invisible', '--convert-to', 'pdf', '--outdir', os.getcwd(), temp_docx], check=True)
        if os.path.exists(temp_pdf): return send_file(temp_pdf, as_attachment=True, download_name='converted.pdf', mimetype='application/pdf')
        else: return "PDF olu≈üturulamadƒ±", 500
    except Exception as e: return f"Sunucu Hatasƒ±: {str(e)}", 500
    finally:
        if os.path.exists(temp_docx): os.remove(temp_docx)
        if os.path.exists(temp_pdf): os.remove(temp_pdf)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
